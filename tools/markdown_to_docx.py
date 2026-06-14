from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


import sys

ROOT = Path(__file__).resolve().parents[1]

if len(sys.argv) >= 2:
    SOURCE = ROOT / sys.argv[1]
    OUTPUT = SOURCE.with_suffix(".docx")
else:
    SOURCE = ROOT / "DOKUMENTIMI_AKADEMIK.md"
    OUTPUT = ROOT / "DOKUMENTIMI_AKADEMIK.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cell, value, bold=row_index == 0)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
            if row_index == 0:
                set_cell_shading(cell, "DDEBE7")

    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.style = "Code Block"
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Cascadia Mono"
        run.font.size = Pt(8.5)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for level, size, color in [
        (1, 18, "134E4A"),
        (2, 14, "0F766E"),
        (3, 12, "18211F"),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    code = styles.add_style("Code Block", 1)
    code.font.name = "Cascadia Mono"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor(32, 43, 42)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_after = Pt(1)
    code.paragraph_format.line_spacing = 1.0


def add_cover_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("QR Event API")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(19, 78, 74)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    run = subtitle.add_run("Dokumentim akademik per REST API, frontend, databaze dhe gjenerim te kodeve QR")
    run.font.name = "Aptos"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(82, 96, 93)

    meta = document.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    data = [
        ("Teknologjia kryesore", "ASP.NET Core Web API"),
        ("Databaza", "SQLite me Entity Framework Core"),
        ("Siguria", "JWT Bearer Authentication"),
        ("Dokumentimi", "Swagger / OpenAPI"),
    ]
    for row_index, (label, value) in enumerate(data):
        cells = meta.rows[row_index].cells
        set_cell_shading(cells[0], "DDEBE7")
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)

    document.add_page_break()


def add_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "QR Event API - Dokumentim Akademik"
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(102, 115, 111)


def parse_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()

        if not line:
            index += 1
            continue

        if line.startswith("# "):
            # Cover already includes the main title.
            index += 1
            continue

        if line.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            table_lines: list[str] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1

            rows = []
            for table_line in table_lines:
                parts = [part.strip() for part in table_line.strip("|").split("|")]
                if all(re.fullmatch(r":?-{2,}:?", part) for part in parts):
                    continue
                rows.append(parts)
            add_table(document, rows)
            continue

        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            index += 1
            continue

        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            index += 1
            continue

        if line.startswith("- "):
            add_bullet(document, line[2:])
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            add_numbered(document, numbered.group(1))
            index += 1
            continue

        add_paragraph(document, line)
        index += 1


def main() -> None:
    document = Document()
    configure_styles(document)
    add_cover_page(document)

    # Keep the body in a fresh section so the report has normal margins after the cover.
    document.add_section(WD_SECTION.NEW_PAGE)
    parse_markdown(document, SOURCE.read_text(encoding="utf-8"))
    add_footer(document)
    document.save(OUTPUT)


if __name__ == "__main__":
    main()
